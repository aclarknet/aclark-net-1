from django.http import HttpResponse
from django_xhtml2pdf.utils import generate_pdf
from docx import Document
from lxml import etree
from openpyxl import Workbook

# from openpyxl.utils import get_column_letter
# from openpyxl.styles import Font
# from openpyxl.styles import PatternFill


def render_doc(context, **kwargs):
    """
    """

    # https://python-docx.readthedocs.io/en/latest/#what-it-can-do
    document = Document()
    filename = kwargs.get("filename")
    item = context["item"]
    document.add_heading(item.title, 0)
    root = etree.fromstring(item.text)
    for elem in root.iter():
        if elem.tag == "h1":
            document.add_heading(elem.text, level=1)
            document.add_paragraph("")
        if elem.tag == "h2":
            document.add_heading(elem.text, level=2)
            document.add_paragraph("")
        if elem.tag == "h3":
            document.add_heading(elem.text, level=3)
            document.add_paragraph("")
        if elem.tag == "h4":
            document.add_heading(elem.text, level=4)
            document.add_paragraph("")
        if elem.tag == "h5":
            document.add_heading(elem.text, level=5)
            document.add_paragraph("")
        if elem.tag == "h6":
            document.add_heading(elem.text, level=6)
            document.add_paragraph("")
        if elem.tag == "p":
            document.add_paragraph(elem.text)
            document.add_paragraph("")

    response = HttpResponse(content_type="docx")
    response["Content-Disposition"] = "attachment; filename=%s" % filename
    document.save(response)
    return response


def render_pdf(context, **kwargs):
    """
    """

    filename = kwargs.get("filename")
    template = kwargs.get("template")
    response = HttpResponse(content_type="application/pdf")
    response["Content-Disposition"] = "filename=%s" % filename
    return generate_pdf(template, context=context, file_object=response)


def render_xls(context, **kwargs):
    """
    """

    # https://openpyxl.readthedocs.io/en/stable/usage.html#write-a-workbook
    workbook = Workbook()
    filename = kwargs.get("filename")
    item = context["item"]
    sheet1 = workbook.active
    sheet1.title = "Invoice"
    for entry in item.time_set.all():
        if entry.task:
            sheet1.append(
                [
                    entry.date,
                    entry.task.name,
                    entry.description,
                    entry.hours,
                    entry.task.rate,
                ]
            )
        else:
            sheet1.append(
                [entry.date, "Task name", entry.description, entry.hours, "Task rate",]
            )
    response = HttpResponse(content_type="xlsx")
    response["Content-Disposition"] = "attachment; filename=%s" % filename
    workbook.save(response)
    return response
